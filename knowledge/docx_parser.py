"""
Парсер Word-документов (.docx) для каталога и скриптов продаж.
"""

import logging
from pathlib import Path
from docx import Document

logger = logging.getLogger(__name__)


def parse_catalog_docx(file_path: str) -> list[dict]:
    """
    Парсит каталог товаров из .docx.
    Разбивает по заголовкам или двойным переносам строк.
    Возвращает список чанков: {'text': str, 'metadata': dict}
    """
    doc = Document(file_path)
    chunks = []
    current_chunk = {"text": "", "metadata": {"source": Path(file_path).name}}
    current_heading = ""

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            # Пустая строка — возможная граница между товарами
            if current_chunk["text"].strip():
                chunks.append(current_chunk)
                current_chunk = {"text": "", "metadata": {"source": Path(file_path).name}}
            continue

        # Заголовок = название товара или раздела
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            if current_chunk["text"].strip():
                chunks.append(current_chunk)
            current_heading = text
            current_chunk = {
                "text": text,
                "metadata": {
                    "source": Path(file_path).name,
                    "product_name": text,
                    "section": current_heading,
                },
            }
        else:
            current_chunk["text"] += "\n" + text
            # Пробуем извлечь цену из текста
            _extract_price(text, current_chunk["metadata"])

    if current_chunk["text"].strip():
        chunks.append(current_chunk)

    # Парсим таблицы (размеры, цены)
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            table_text = "\n".join([" | ".join(row) for row in rows])
            chunks.append({
                "text": table_text,
                "metadata": {"source": Path(file_path).name, "type": "table"},
            })

    logger.info(f"Parsed {file_path}: {len(chunks)} chunks")
    return chunks


def parse_scripts_docx(file_path: str) -> list[dict]:
    """
    Парсит скрипты продаж из .docx.
    Разбивает по заголовкам/секциям.
    """
    doc = Document(file_path)
    chunks = []
    current_text = ""
    current_heading = "Общее"

    for para in doc.paragraphs:
        text = para.text.strip()
        if para.style and para.style.name.startswith("Heading"):
            if current_text.strip():
                chunks.append({
                    "text": current_text.strip(),
                    "metadata": {
                        "source": Path(file_path).name,
                        "section": current_heading,
                        "type": "sales_script",
                    },
                })
            current_heading = text
            current_text = text + "\n"
        elif text:
            current_text += text + "\n"

    if current_text.strip():
        chunks.append({
            "text": current_text.strip(),
            "metadata": {
                "source": Path(file_path).name,
                "section": current_heading,
                "type": "sales_script",
            },
        })

    logger.info(f"Parsed scripts {file_path}: {len(chunks)} chunks")
    return chunks


def _extract_price(text: str, metadata: dict):
    """Пробует извлечь цену из строки текста."""
    import re
    # Паттерны: "Цена: 15 000 тг", "45000тг", "15 000 ₸"
    patterns = [
        r"[Цц]ена[:\s]*(\d[\d\s]*)\s*(?:тг|тенге|₸|руб|р\.?)",
        r"(\d[\d\s]*)\s*(?:тг|тенге|₸)",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            price = match.group(1).replace(" ", "")
            metadata["price"] = price
            break
